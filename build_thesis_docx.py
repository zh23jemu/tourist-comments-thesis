from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\Coding\tourist-comments-thesis")
TEMPLATE = ROOT / r"2026届毕业论文模板\人文社科类\1.毕业论文模板 .docx"
SOURCE = ROOT / "本科毕业论文初稿-大数据背景下旅游评论挖掘及景区推荐研究.md"
OUTPUT = ROOT / "本科毕业论文-正式排版稿.docx"


TITLE = "大数据背景下旅游评论挖掘及景区推荐研究"
SUBTITLE = "以去哪儿为例"
LEVEL3_COUNTER = {}


def set_run_font(run, name: str, size: int, bold: bool = False) -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_style(paragraph, style_name: str | None) -> None:
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            pass


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    font_name: str = "宋体",
    font_size: int = 12,
    bold: bool = False,
    first_line_chars: int | None = None,
    line_spacing: float | None = 1.25,
    space_before: int = 0,
    space_after: int = 0,
    style_name: str | None = None,
):
    p = doc.add_paragraph()
    set_paragraph_style(p, style_name)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if first_line_chars is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_chars * 10)
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold=bold)
    return p


def clear_paragraph(paragraph) -> None:
    """清空模板段落中的占位内容，但保留段落属性、分节符和样式。"""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_font(paragraph, font_name: str, font_size: int, bold: bool = False) -> None:
    """统一设置段落所有文字的中英文字体，避免模板占位格式残留。"""
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size, bold=bold)


def replace_paragraph_text(
    paragraph,
    text: str = "",
    *,
    font_name: str = "宋体",
    font_size: int = 12,
    bold: bool = False,
    align=None,
    first_line_chars: int | None = None,
    line_spacing: float | None = None,
    space_before: int | None = None,
    space_after: int | None = None,
    style_name: str | None = None,
):
    """在模板原有段落位置写入正文，最大限度继承模板的分页和分节结构。"""
    clear_paragraph(paragraph)
    set_paragraph_style(paragraph, style_name)
    if align is not None:
        paragraph.alignment = align
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line_chars is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_chars * 12)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size, bold=bold)
    return paragraph


def append_paragraph_after(
    doc: Document,
    anchor,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    font_name: str = "宋体",
    font_size: int = 12,
    bold: bool = False,
    first_line_chars: int | None = None,
    line_spacing: float | None = 1.25,
    space_before: int = 0,
    space_after: int = 0,
    style_name: str | None = None,
):
    """在指定段落后插入新段落，用于在保留模板前置结构时继续写正文。"""
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    p = doc.paragraphs[-1].__class__(new_p, anchor._parent)
    set_paragraph_style(p, style_name)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if first_line_chars is not None:
        p.paragraph_format.first_line_indent = Pt(first_line_chars * 12)
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold=bold)
    return p


def insert_section_after(anchor, section_index: int) -> None:
    """复用模板中的分节设置，确保新生成内容的页眉页脚、页边距与模板一致。"""
    template = Document(TEMPLATE)
    sect_pr = deepcopy(template.sections[section_index]._sectPr)
    p_pr = anchor._p.get_or_add_pPr()
    old = p_pr.sectPr
    if old is not None:
        p_pr.remove(old)
    p_pr.append(sect_pr)


def remove_trailing_template_content(doc: Document, keep_count: int) -> None:
    """删除模板正文占位说明，只保留封面、承诺书、授权书、摘要和目录的骨架。"""
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body)[keep_count:]:
        if child is not sect_pr:
            body.remove(child)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_char_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_char_sep)

    end_run = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char_end)


def enable_update_fields(doc: Document) -> None:
    """让 Word 打开文档时自动更新目录等域，避免目录停留在模板占位内容。"""
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def add_tab_stop(p, pos_twips: int) -> None:
    p_pr = p._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)


def parse_markdown_sections(text: str):
    lines = text.splitlines()
    sections = []
    current = None
    buffer: list[str] = []

    def flush():
        nonlocal buffer, current
        if current is not None:
            current["content"] = "\n".join(buffer).strip()
            sections.append(current)
        buffer = []

    for line in lines:
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            flush()
            current = {"level": 2, "title": line[3:].strip(), "content": ""}
            continue
        if line.startswith("### "):
            flush()
            current = {"level": 3, "title": line[4:].strip(), "content": ""}
            continue
        if line.startswith("#### "):
            flush()
            current = {"level": 4, "title": line[5:].strip(), "content": ""}
            continue
        buffer.append(line)
    flush()
    return sections


def split_paragraphs(content: str) -> list[str]:
    parts = re.split(r"\n\s*\n", content.strip())
    cleaned = []
    for p in parts:
        text = p.strip()
        if not text:
            continue
        if text.startswith("**关键词：") or text.startswith("**Key words:"):
            continue
        cleaned.append(text)
    return cleaned


def write_cover(doc: Document) -> None:
    """替换模板封面占位文字，同时保留封面原有版式。"""
    paragraphs = doc.paragraphs
    replace_paragraph_text(paragraphs[3], "本科毕业论文", font_name="黑体", font_size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph_text(paragraphs[4], TITLE, font_name="黑体", font_size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph_text(paragraphs[5], SUBTITLE, font_name="黑体", font_size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, label in zip(range(8, 14), ["学    院", "专    业", "班    级", "学    生", "学    号", "指导教师"]):
        replace_paragraph_text(
            paragraphs[idx],
            f"{label}      ____________________",
            font_name="宋体",
            font_size=16,
            bold=False,
            line_spacing=1.25,
        )
        paragraphs[idx].paragraph_format.first_line_indent = Pt(48.2)

    replace_paragraph_text(paragraphs[16], "二〇二六年四月", font_name="宋体", font_size=16, align=WD_ALIGN_PARAGRAPH.CENTER)


def write_abstracts(doc: Document, sections, anchors) -> None:
    abstract = next(s for s in sections if s["title"] == "摘要")
    english = next(s for s in sections if s["title"] == "Abstract")

    replace_paragraph_text(
        anchors["zh_abstract_title"],
        "摘    要",
        font_name="黑体",
        font_size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        space_before=31.2,
        space_after=14,
        style_name="中文摘要",
    )
    cursor = anchors["zh_abstract_title"]
    for para in split_paragraphs(abstract["content"]):
        cursor = append_paragraph_after(doc, cursor, para, first_line_chars=2, line_spacing=1.25)

    cursor = append_paragraph_after(doc, cursor, "", font_name="宋体", font_size=10, line_spacing=1.25, space_before=0, space_after=0)
    r1 = cursor.add_run("关键词：")
    set_run_font(r1, "黑体", 10, True)
    r2 = cursor.add_run("大数据；旅游评论；文本挖掘；景区推荐；情感分析；去哪儿")
    set_run_font(r2, "宋体", 10, bold=False)
    cursor.paragraph_format.first_line_indent = Pt(24)

    replace_paragraph_text(
        anchors["en_abstract_title"],
        "ABSTRACT",
        font_name="Times New Roman",
        font_size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        space_before=31.2,
        space_after=14,
        style_name="英文摘要",
    )
    cursor = anchors["en_abstract_title"]
    for para in split_paragraphs(english["content"]):
        cursor = append_paragraph_after(doc, cursor, para, font_name="Times New Roman", first_line_chars=2, line_spacing=1.25)

    cursor = append_paragraph_after(doc, cursor, "", font_name="Times New Roman", font_size=10, line_spacing=1.25, space_before=0, space_after=0)
    cursor.paragraph_format.first_line_indent = Pt(24)
    r1 = cursor.add_run("Key words: ")
    set_run_font(r1, "Times New Roman", 10, True)
    r2 = cursor.add_run("big data; tourism reviews; text mining; scenic spot recommendation; sentiment analysis; Qunar")
    set_run_font(r2, "Times New Roman", 10, bold=False)


def write_toc(doc: Document, anchors) -> None:
    replace_paragraph_text(anchors["toc_title"], "目    录", font_name="黑体", font_size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=31.2, space_after=14)
    p = anchors["toc_field"]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(p, r'TOC \o "1-3" \h \z \u')


def normalize_heading(title: str) -> str:
    mapping = {
        "一、绪论": "一、绪论",
        "二、提出问题：旅游评论挖掘与景区推荐的理论基础": "二、旅游评论挖掘与景区推荐的理论基础",
        "三、分析问题：研究设计与数据处理": "三、研究设计与数据处理",
        "四、分析问题：旅游评论挖掘实证分析": "四、旅游评论挖掘实证分析",
        "五、分析问题：景区推荐模型构建与结果分析": "五、景区推荐模型构建与结果分析",
        "六、解决问题：系统设计、应用价值与优化路径": "六、系统设计、应用价值与优化路径",
        "七、结论与展望": "七、结论与展望",
        "参考文献": "参考文献",
    }
    return mapping.get(title, title)


CN_NUMS = "一二三四五六七八九十"


def to_cn_index(index: int) -> str:
    if 1 <= index <= 10:
        return CN_NUMS[index - 1]
    return str(index)


def format_level3(title: str, idx: int) -> str:
    m = re.match(r"^(\d+\.\d+)\s*(.*)$", title.strip())
    raw = m.group(2).strip() if m else title.strip()
    raw = re.sub(r"^[（(][一二三四五六七八九十]+[）)]\s*", "", raw)
    return f"（{to_cn_index(idx)}）{raw}"


def format_level4(title: str, idx: int) -> str:
    raw = re.sub(r"^\d+\.\d+\.\d+\s*", "", title).strip()
    raw = re.sub(r"^[（(][一二三四五六七八九十]+[）)]\s*", "", raw)
    raw = re.sub(r"^\d+[.．、]\s*", "", raw)
    return f"{idx}. {raw}"


def write_body(doc: Document, sections, anchor) -> None:
    insert_section_after(anchor, 4)
    cursor = anchor
    current_level2 = 0
    current_level3 = 0
    current_level4 = 0

    for sec in sections:
        title = sec["title"]
        if title in {"摘要", "Abstract", "以去哪儿为例"}:
            continue

        if sec["level"] == 2:
            if title in {"七、结论与展望", "参考文献"}:
                insert_section_after(cursor, 5)
            current_level2 += 1
            current_level3 = 0
            current_level4 = 0
            cursor = append_paragraph_after(
                doc,
                cursor,
                normalize_heading(title),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                font_name="黑体",
                font_size=15,
                bold=True,
                line_spacing=1.5,
                space_before=0,
                space_after=14,
                style_name="章的标题",
            )
            if title == "参考文献":
                for para in split_paragraphs(sec["content"]):
                    cursor = append_paragraph_after(doc, cursor, para, font_name="宋体", font_size=12, first_line_chars=0, line_spacing=1.25, style_name="论文正文")
                continue

        elif sec["level"] == 3:
            current_level3 += 1
            current_level4 = 0
            heading = format_level3(title, current_level3)
            cursor = append_paragraph_after(
                doc,
                cursor,
                heading,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_name="黑体",
                font_size=14,
                bold=True,
                line_spacing=1.5,
                space_before=8,
                space_after=0,
                style_name="节标题",
            )
        elif sec["level"] == 4:
            current_level4 += 1
            heading = format_level4(title, current_level4)
            cursor = append_paragraph_after(
                doc,
                cursor,
                heading,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_name="黑体",
                font_size=12,
                bold=True,
                line_spacing=1.5,
                space_before=8,
                space_after=0,
                style_name="段标题",
            )

        for para in split_paragraphs(sec["content"]):
            clean = para.replace("**", "")
            cursor = append_paragraph_after(doc, cursor, clean, first_line_chars=2, style_name="论文正文")


def prepare_template_body(doc: Document):
    """保留模板前置固定页，删除正文格式说明占位内容，并清理摘要/目录页说明文字。"""
    from docx.text.paragraph import Paragraph

    anchors = {
        "zh_abstract_title": doc.paragraphs[37],
        "en_abstract_title": doc.paragraphs[47],
        "toc_title": doc.paragraphs[59],
        "toc_field": doc.paragraphs[60],
    }

    # 先清空摘要、英文摘要和目录页中模板说明，保留关键分节段落。
    for idx in list(range(38, 47)) + list(range(48, 59)) + list(range(61, 85)):
        clear_paragraph(doc.paragraphs[idx])

    # 截断模板正文说明，只保留到目录页起始骨架。
    body = doc._element.body
    children = list(body)
    keep_count = 85
    anchor_element = children[84]
    remove_trailing_template_content(doc, keep_count)

    anchors["body_anchor"] = Paragraph(anchor_element, doc._body)
    return anchors


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    sections = parse_markdown_sections(text)

    doc = Document(TEMPLATE)

    write_cover(doc)
    anchors = prepare_template_body(doc)
    write_abstracts(doc, sections, anchors)
    write_toc(doc, anchors)
    write_body(doc, sections, anchors["body_anchor"])
    enable_update_fields(doc)

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    main()
